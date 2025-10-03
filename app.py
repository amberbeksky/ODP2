import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import traceback
from tkcalendar import DateEntry
from datetime import datetime, timedelta
import gspread
from google.oauth2.service_account import Credentials
import os
import json
import sys
import updater
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import simpledialog
import time

# ================== Пути ==================
APP_DIR = os.path.join(os.getenv("APPDATA") or os.path.expanduser("~"), "MyApp")
os.makedirs(APP_DIR, exist_ok=True)

DB_NAME = os.path.join(APP_DIR, "clients.db")
SHEET_ID = "1_DfTT8yzCjP0VH0PZu1Fz6FYMm1eRr7c0TmZU2DrH_w"

# ================== ИМПОРТ МЕНЕДЖЕРА АУТЕНТИФИКАЦИИ ==================
try:
    from auth_manager import AuthManager
    AUTH_AVAILABLE = True
except ImportError:
    AUTH_AVAILABLE = False
    print("⚠️ Модуль auth_manager не найден. Аутентификация отключена.")

# ================== Менеджер настроек ==================
class SettingsManager:
    def __init__(self):
        self.settings_file = os.path.join(APP_DIR, "settings.json")
        self.settings = self.load_settings()
    
    def load_settings(self):
        """Загрузка настроек из файла"""
        default_settings = {
            'default_export_path': os.path.join(os.path.expanduser("~"), "Desktop"),
            'auto_check_updates': True,
            'show_notifications': True,
            'theme': 'modern'
        }
        
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    loaded_settings = json.load(f)
                    default_settings.update(loaded_settings)
        except Exception as e:
            print(f"Ошибка загрузки настроек: {e}")
        
        return default_settings
    
    def save_settings(self):
        """Сохранение настроек в файл"""
        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Ошибка сохранения настроек: {e}")
    
    def get(self, key, default=None):
        """Получить значение настройки"""
        return self.settings.get(key, default)
    
    def set(self, key, value):
        """Установить значение настройки"""
        self.settings[key] = value
        self.save_settings()

# Глобальный экземпляр менеджера настроек
settings_manager = SettingsManager()

# Глобальный экземпляр менеджера аутентификации
auth_manager = None

# ================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ==================
def create_tooltip(widget, text):
    """Создание всплывающей подсказки для виджета"""
    def on_enter(event):
        tooltip = tk.Toplevel()
        tooltip.wm_overrideredirect(True)
        tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
        
        label = tk.Label(tooltip, text=text, background="#ffffe0", 
                        relief='solid', borderwidth=1, font=ModernStyle.FONTS['small'])
        label.pack()
        
        widget.tooltip = tooltip
    
    def on_leave(event):
        if hasattr(widget, 'tooltip') and widget.tooltip:
            widget.tooltip.destroy()
            widget.tooltip = None
    
    widget.bind("<Enter>", on_enter)
    widget.bind("<Leave>", on_leave)

def show_notifications():
    """Показать уведомления (вызывается из меню)"""
    notification_system.show_notification_window()

def show_statistics():
    """Показать статистику по клиентам"""
    clients = get_all_clients(limit=10000)
    total = len(clients)
    
    today = datetime.today().date()
    active = 0
    expired = 0
    soon = 0
    groups = {}
    
    for client in clients:
        ippcu_end = client[8]
        group = client[9] or "Без группы"
        
        if group not in groups:
            groups[group] = 0
        groups[group] += 1
        
        if ippcu_end:
            try:
                end_date = datetime.strptime(ippcu_end, "%Y-%m-%d").date()
                if end_date < today:
                    expired += 1
                elif end_date <= today + timedelta(days=30):
                    soon += 1
                else:
                    active += 1
            except:
                pass
    
    stats_text = f"""📊 СТАТИСТИКА

Всего клиентов: {total}
├─ Активные ИППСУ: {active}
├─ Истекают в течение 30 дней: {soon}
└─ Просроченные ИППСУ: {expired}

📂 РАСПРЕДЕЛЕНИЕ ПО ГРУППАМ:"""
    
    for group, count in sorted(groups.items()):
        percentage = (count / total) * 100 if total > 0 else 0
        stats_text += f"\n├─ {group}: {count} чел. ({percentage:.1f}%)"
    
    messagebox.showinfo("📊 Статистика", stats_text)

def check_expiring_ippcu():
    """Проверка истекающих ИППСУ при запуске"""
    clients = get_all_clients(limit=10000)
    today = datetime.today().date()
    
    expiring = []
    expired = []
    
    for client in clients:
        ippcu_end = client[8]
        if ippcu_end:
            try:
                end_date = datetime.strptime(ippcu_end, "%Y-%m-%d").date()
                days_left = (end_date - today).days
                
                if 0 <= days_left <= 7:
                    expiring.append((client, days_left))
                elif days_left < 0:
                    expired.append((client, abs(days_left)))
            except:
                pass
    
    messages = []
    
    if expired:
        messages.append(f"❌ ПРОСРОЧЕНЫ {len(expired)} ИППСУ!")
        for client, days in expired[:3]:
            messages.append(f"   {client[1]} {client[2]} - просрочено {days} дн. назад")
    
    if expiring:
        messages.append(f"⚠️ ИСТЕКАЮТ {len(expiring)} ИППСУ в течение недели!")
        for client, days in expiring[:3]:
            messages.append(f"   {client[1]} {client[2]} - осталось {days} дн.")
    
    if messages:
        messagebox.showwarning("Внимание!", "\n".join(messages))

def export_selected_to_word():
    selected_items = []
    for row_id in tree.get_children():
        values = tree.item(row_id, "values")
        if values and values[0] == "X":
            selected_items.append(values)

    if not selected_items:
        messagebox.showerror("Ошибка", "Отметьте галочками хотя бы одного клиента")
        return

    shift_name = simpledialog.askstring("Смена", "Введите название смены (например: 11 смена)")
    if not shift_name:
        return

    date_range = simpledialog.askstring("Даты", "Введите период (например: с 01.10.2024 по 15.10.2024)")
    if not date_range:
        return

    doc = Document()

    heading = doc.add_paragraph(f"{shift_name} {date_range}")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    for i, values in enumerate(selected_items, start=1):
        last = values[2]
        first = values[3]
        middle = values[4]
        dob = values[5]

        fio = " ".join(v for v in [last, first, middle] if v)
        p = doc.add_paragraph(f"{i}. {fio} – {dob} г.р.")
        p.runs[0].font.size = Pt(12)

    spacer = doc.add_paragraph("\n")
    spacer.paragraph_format.space_after = Pt(300)

    total = len(selected_items)
    total_p = doc.add_paragraph(f"Итого: {total} человек")
    total_p.runs[0].bold = True
    total_p.runs[0].font.size = Pt(12)

    podpis = doc.add_paragraph()
    podpis.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run_role = podpis.add_run("Заведующая отделением дневного пребывания ")
    run_role.font.size = Pt(12)

    run_line = podpis.add_run("__________________ ")
    run_line.font.size = Pt(12)

    run_name = podpis.add_run("Дурандина А.В.")
    run_name.font.size = Pt(12)

    # Используем путь из настроек или рабочий стол по умолчанию
    export_path = settings_manager.get('default_export_path', os.path.join(os.path.expanduser("~"), "Desktop"))
    
    safe_shift = shift_name.replace(" ", "_")
    safe_date = date_range.replace(" ", "_").replace(":", "-").replace(".", "-")
    file_name = f"{safe_shift}_{safe_date}.docx"
    file_path = os.path.join(export_path, file_name)

    try:
        doc.save(file_path)
        messagebox.showinfo("Готово", f"Список сохранён:\n{file_path}")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")

    for row_id in tree.get_children():
        values = list(tree.item(row_id, "values"))
        if values[0] == "X":
            values[0] = " "
            tree.item(row_id, values=values)
    
    if hasattr(root, 'update_word_count'):
        root.update_word_count()

def show_status_message(message, duration=3000):
    """Показать временное сообщение в статусной строке"""
    if hasattr(root, 'status_label'):
        root.status_label.config(text=message)
        root.after(duration, lambda: root.status_label.config(text="Готово"))

def copy_to_clipboard(text):
    """Копировать текст в буфер обмена"""
    if text:
        root.clipboard_clear()
        root.clipboard_append(text)
        show_status_message(f"Скопировано: {text[:20]}..." if len(text) > 20 else f"Скопировано: {text}")

def add_to_word_list(item):
    """Добавить/убрать клиента из списка для Word"""
    values = list(tree.item(item, "values"))
    values[0] = "X" if values[0].strip() == "" else " "
    tree.item(item, values=values)
    
    action = "добавлен в" if values[0] == "X" else "удален из"
    show_status_message(f"Клиент {action} списка для Word")

def quick_view(client_id):
    """Быстрый просмотр информации о клиенте"""
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        cur.execute(
            "SELECT last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name FROM clients WHERE id=?",
            (client_id,)
        )
        client = cur.fetchone()
    
    if not client:
        messagebox.showerror("Ошибка", "Клиент не найден")
        return
    
    last, first, middle, dob, phone, contract, ippcu_start, ippcu_end, group = client
    
    info_text = f"""👤 {last} {first} {middle or ''}

📅 Дата рождения: {dob or 'не указана'}
📞 Телефон: {phone or 'не указан'}
📄 Договор: {contract or 'не указан'}
🏷️ Группа: {group or 'не указана'}

📋 ИППСУ:
   Начало: {ippcu_start or 'не указано'}
   Окончание: {ippcu_end or 'не указано'}"""
    
    if ippcu_end:
        try:
            end_date = datetime.strptime(ippcu_end, "%Y-%m-%d").date()
            today = datetime.today().date()
            days_left = (end_date - today).days
            
            if days_left < 0:
                info_text += f"\n\n⚠️ ИППСУ ПРОСРОЧЕН на {abs(days_left)} дн."
            elif days_left <= 30:
                info_text += f"\n\n⚠️ ИППСУ истекает через {days_left} дн."
            else:
                info_text += f"\n\n✅ ИППСУ активен ({days_left} дн. осталось)"
        except:
            pass
    
    messagebox.showinfo("Информация о клиенте", info_text)

# ================== ФУНКЦИИ АУТЕНТИФИКАЦИИ ==================
def check_auth_status():
    """Проверка статуса авторизации"""
    if not AUTH_AVAILABLE:
        return True  # Пропускаем проверку если модуль недоступен
    
    if auth_manager and hasattr(auth_manager, 'current_user') and auth_manager.current_user:
        return True
    return False

def require_auth(func):
    """Декоратор для проверки авторизации"""
    def wrapper(*args, **kwargs):
        if not check_auth_status():
            show_login_window()
            return None
        return func(*args, **kwargs)
    return wrapper

def show_login_window():
    """Окно входа в систему"""
    login_window = tk.Toplevel(root)
    login_window.title("Авторизация - Отделение дневного пребывания")
    login_window.geometry("450x450")
    login_window.configure(bg=ModernStyle.COLORS['background'])
    login_window.resizable(False, False)
    
    # Центрирование окна
    login_window.transient(root)
    login_window.grab_set()
    
    # Заголовок
    header = tk.Frame(login_window, bg=ModernStyle.COLORS['primary'], height=80)
    header.pack(fill='x', padx=0, pady=0)
    
    tk.Label(header, text="🔐 Авторизация", 
            bg=ModernStyle.COLORS['primary'],
            fg='white',
            font=ModernStyle.FONTS['h1']).pack(pady=20)
    
    # Основное содержимое
    content_frame = tk.Frame(login_window, bg=ModernStyle.COLORS['background'], padx=30, pady=30)
    content_frame.pack(fill='both', expand=True)
    
    # Поля ввода
    fields_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    fields_frame.pack(fill='both', expand=True, pady=20)
    
    # Логин
    login_frame = tk.Frame(fields_frame, bg=ModernStyle.COLORS['background'])
    login_frame.pack(fill='x', pady=10)
    
    tk.Label(login_frame, text="Логин:",
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_primary'],
            font=ModernStyle.FONTS['body']).pack(anchor='w')
    
    login_var = tk.StringVar()
    login_entry = tk.Entry(login_frame, textvariable=login_var,
                          font=ModernStyle.FONTS['body'], width=30)
    login_entry.pack(fill='x', pady=5)
    
    # Пароль
    password_frame = tk.Frame(fields_frame, bg=ModernStyle.COLORS['background'])
    password_frame.pack(fill='x', pady=10)
    
    tk.Label(password_frame, text="Пароль:",
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_primary'],
            font=ModernStyle.FONTS['body']).pack(anchor='w')
    
    password_var = tk.StringVar()
    password_entry = tk.Entry(password_frame, textvariable=password_var,
                             show="•", font=ModernStyle.FONTS['body'], width=30)
    password_entry.pack(fill='x', pady=5)
    
    # Чекбокс "Запомнить меня"
    remember_var = tk.BooleanVar(value=True)
    remember_frame = tk.Frame(fields_frame, bg=ModernStyle.COLORS['background'])
    remember_frame.pack(fill='x', pady=10)
    
    remember_cb = tk.Checkbutton(remember_frame, 
                                text="Запомнить меня на 30 дней",
                                variable=remember_var,
                                bg=ModernStyle.COLORS['background'],
                                fg=ModernStyle.COLORS['text_primary'],
                                font=ModernStyle.FONTS['small'],
                                selectcolor=ModernStyle.COLORS['primary'])
    remember_cb.pack(anchor='w')
    
    # Информация о пользователях
    info_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['surface'],
                         relief='solid', bd=1, padx=15, pady=10)
    info_frame.pack(fill='x', pady=10)
    
    info_text = """Доступные пользователи:
• admin / admin - Зеленков Д.В. (Администратор)
• ДУРАНДИНА / 12345 - Дурандина А.В. (Заведующая)
• ЛАВРОВА / 12345 - Лаврова А.А. (Сотрудник)"""
    
    tk.Label(info_frame, text=info_text,
            bg=ModernStyle.COLORS['surface'],
            fg=ModernStyle.COLORS['text_secondary'],
            font=ModernStyle.FONTS['small'],
            justify='left').pack(anchor='w')
    
    # Кнопки
    button_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    button_frame.pack(fill='x', pady=10)
    
    def attempt_login():
        username = login_var.get().strip()
        password = password_var.get()
        
        print(f"DEBUG: Attempting login for {username}")  # ДЕБАГ
        
        if not username or not password:
            messagebox.showerror("Ошибка", "Введите логин и пароль")
            return
        
        try:
            success, message = auth_manager.login(username, password, remember_var.get())
            
            if success:
                print("DEBUG: Login successful!")  # ДЕБАГ
                login_window.destroy()
                # Очищаем главное окно и создаем интерфейс заново
                for widget in root.winfo_children():
                    widget.destroy()
                initialize_main_application()
                show_status_message(f"Добро пожаловать, {auth_manager.get_user_display_name()}!")
            else:
                print(f"DEBUG: Login failed: {message}")  # ДЕБАГ
                # Показываем конкретную ошибку
                if "locked" in message.lower():
                    retry = messagebox.askretrycancel(
                        "Ошибка базы данных", 
                        "База данных временно заблокирована. Повторить попытку?"
                    )
                    if retry:
                        root.after(1000, attempt_login)  # Повторить через 1 секунду
                else:
                    messagebox.showerror("Ошибка входа", message)
                password_var.set("")
                password_entry.focus()
        except sqlite3.OperationalError as e:
            if "locked" in str(e):
                retry = messagebox.askretrycancel(
                    "Ошибка базы данных", 
                    "База данных заблокирована. Закройте другие экземпляры программы и попробуйте снова."
                )
                if retry:
                    root.after(1000, attempt_login)
            else:
                messagebox.showerror("Ошибка", f"Ошибка при входе: {e}")
        except Exception as e:
            print(f"DEBUG: Login exception: {e}")  # ДЕБАГ
            messagebox.showerror("Ошибка", f"Ошибка при входе: {e}")
    
    # Кнопка Войти
    login_btn = ttk.Button(button_frame, text="Войти", 
                          style='Primary.TButton',
                          command=attempt_login)
    login_btn.pack(fill='x', pady=5)
    
    # Кнопка Отмена
    def cancel_login():
        login_window.destroy()
        if not check_auth_status():
            root.destroy()
    
    cancel_btn = ttk.Button(button_frame, text="Отмена", 
                           style='Secondary.TButton',
                           command=cancel_login)
    cancel_btn.pack(fill='x', pady=5)
    
    # Обработка нажатия Enter
    def on_enter_pressed(event):
        attempt_login()
    
    login_entry.bind('<Return>', on_enter_pressed)
    password_entry.bind('<Return>', on_enter_pressed)
    
    # Фокус на поле логина
    login_entry.focus()
    
    # Если есть запомненный пользователь, закрываем окно входа и инициализируем приложение
    if auth_manager and hasattr(auth_manager, 'current_user') and auth_manager.current_user:
        login_window.destroy()
        initialize_main_application()

def show_user_profile():
    """Окно профиля пользователя"""
    if not auth_manager or not auth_manager.current_user:
        messagebox.showinfo("Информация", "Вы не авторизованы")
        return
    
    profile_window = tk.Toplevel(root)
    profile_window.title("Профиль пользователя")
    profile_window.geometry("450x350")
    profile_window.configure(bg=ModernStyle.COLORS['background'])
    profile_window.resizable(False, False)
    
    # Заголовок
    header = tk.Frame(profile_window, bg=ModernStyle.COLORS['primary'], height=60)
    header.pack(fill='x', padx=0, pady=0)
    
    tk.Label(header, text="👤 Профиль", 
            bg=ModernStyle.COLORS['primary'],
            fg='white',
            font=ModernStyle.FONTS['h2']).pack(pady=15)
    
    # Основное содержимое
    content_frame = tk.Frame(profile_window, bg=ModernStyle.COLORS['background'], padx=20, pady=20)
    content_frame.pack(fill='both', expand=True)
    
    user_info = auth_manager.current_user
    
    # Информация о пользователе
    info_text = f"""ФИО: {user_info['full_name']}
Должность: {user_info['role']}
Логин: {user_info['username']}
Права доступа: {', '.join(user_info['permissions'])}
Статус входа: {"Запомнен на 30 дней" if auth_manager.remember_me else "Требуется вход при запуске"}"""
    
    tk.Label(content_frame, text=info_text,
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_primary'],
            font=ModernStyle.FONTS['body'],
            justify='left').pack(anchor='w', pady=10)
    
    # Кнопки
    button_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    button_frame.pack(fill='x', pady=20)
    
    def logout():
        if messagebox.askyesno("Выход", "Вы уверены, что хотите выйти?"):
            auth_manager.logout()
            profile_window.destroy()
            show_login_window()
            update_ui_for_user()
    
    def clear_remember():
        if messagebox.askyesno("Очистка", "Очистить запомненный вход?\nПри следующем запуске потребуется ввод логина и пароля."):
            auth_manager.clear_remember_token()
            auth_manager.remember_me = False
            profile_window.destroy()
            show_status_message("Запомненный вход очищен")
    
    ttk.Button(button_frame, text="Сменить пользователя", 
              style='Primary.TButton',
              command=logout).pack(side='right', padx=(10, 0))
    
    ttk.Button(button_frame, text="Очистить запомненный вход", 
              style='Secondary.TButton',
              command=clear_remember).pack(side='right', padx=(10, 0))
    
    ttk.Button(button_frame, text="Закрыть", 
              style='Secondary.TButton',
              command=profile_window.destroy).pack(side='right')

def update_ui_for_user():
    """Обновление интерфейса в зависимости от прав пользователя"""
    if not auth_manager or not auth_manager.current_user:
        return
    
    # Обновляем заголовок окна
    root.title(f"Отделение дневного пребывания - {auth_manager.get_user_display_name()}")
    
    # Обновляем статусную строку
    if hasattr(root, 'user_status_label'):
        root.user_status_label.config(
            text=f"Пользователь: {auth_manager.get_user_display_name()}"
        )
    
    # Показываем/скрываем элементы в зависимости от прав
    update_permissions_ui()

def update_permissions_ui():
    """Обновление видимости элементов по правам доступа"""
    if not auth_manager or not auth_manager.current_user:
        return
    
    # Пример ограничения функционала для разных ролей
    if hasattr(root, 'add_btn'):
        root.add_btn['state'] = 'normal' if auth_manager.has_permission('edit') else 'disabled'
    
    if hasattr(root, 'delete_btn'):
        root.delete_btn['state'] = 'normal' if auth_manager.has_permission('delete') else 'disabled'

def setup_auth_system():
    """Настройка системы аутентификации"""
    global auth_manager
    
    if not AUTH_AVAILABLE:
        print("⚠️ Аутентификация отключена - модуль auth_manager не найден")
        # Создаем заглушку для случая отсутствия модуля аутентификации
        class AuthStub:
            def __init__(self):
                self.current_user = {
                    'full_name': 'Демо-пользователь',
                    'role': 'Сотрудник',
                    'username': 'demo',
                    'permissions': ['basic']
                }
                self.remember_me = False
            
            def get_user_display_name(self):
                return "Демо-пользователь"
            
            def has_permission(self, permission):
                return True
            
            def logout(self):
                pass
        
        auth_manager = AuthStub()
        return
    
    try:
        auth_manager = AuthManager(DB_NAME)
        # Очищаем просроченные токены при запуске
        auth_manager.cleanup_expired_tokens()
        print("✅ Система аутентификации инициализирована")
    except Exception as e:
        print(f"❌ Ошибка инициализации аутентификации: {e}")
        messagebox.showerror("Ошибка", f"Не удалось инициализировать систему аутентификации: {e}")

# ================== СИСТЕМА УВЕДОМЛЕНИЙ ==================
class NotificationSystem:
    def __init__(self, db_path):
        self.db_path = db_path
        self.notifications = []
        self.is_initialized = False
        
    def initialize(self):
        """Инициализация системы уведомлений (вызывается после инициализации БД)"""
        try:
            # Проверяем, что таблица существует
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.cursor()
                cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='clients'")
                if not cur.fetchone():
                    print("Таблица clients не найдена, уведомления отключены")
                    return False
                    
            self.setup_daily_checks()
            self.is_initialized = True
            return True
        except Exception as e:
            print(f"Ошибка инициализации системы уведомлений: {e}")
            return False
    
    def setup_daily_checks(self):
        """Настройка ежедневных проверок"""
        self.check_birthdays()
        self.check_ippcu_expiry()
        self.check_empty_contracts()
        self.check_upcoming_reviews()
    
    def check_birthdays(self):
        """Проверка ближайших дней рождений"""
        try:
            today = datetime.today().date()
            next_week = today + timedelta(days=30)
            
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.cursor()
                cur.execute("""
                    SELECT last_name, first_name, middle_name, dob 
                    FROM clients 
                    WHERE dob IS NOT NULL AND dob != ''
                """)
                
                clients = cur.fetchall()
            
            for last, first, middle, dob in clients:
                try:
                    if not dob:
                        continue
                        
                    bday = datetime.strptime(dob, "%Y-%m-%d").date()
                    bday_this_year = bday.replace(year=today.year)
                    
                    # Если день рождения уже прошел в этом году, смотрим на следующий год
                    if bday_this_year < today:
                        bday_this_year = bday.replace(year=today.year + 1)
                    
                    days_until = (bday_this_year - today).days
                    
                    if 0 <= days_until <= 30:
                        level = "warning" if days_until <= 7 else "info"
                        self.add_notification(
                            "birthday", 
                            f"🎂 День рождения у {last} {first} {middle or ''} через {days_until} дн. ({bday.strftime('%d.%m.%Y')})",
                            level
                        )
                except ValueError:
                    continue
                    
        except Exception as e:
            print(f"Ошибка проверки дней рождений: {e}")
    
    def check_ippcu_expiry(self):
        """Проверка истекающих ИППСУ"""
        try:
            today = datetime.today().date()
            
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.cursor()
                cur.execute("""
                    SELECT last_name, first_name, ippcu_end 
                    FROM clients 
                    WHERE ippcu_end IS NOT NULL AND ippcu_end != ''
                """)
                
                clients = cur.fetchall()
            
            for last, first, ippcu_end in clients:
                try:
                    if not ippcu_end:
                        continue
                        
                    end_date = datetime.strptime(ippcu_end, "%Y-%m-%d").date()
                    days_left = (end_date - today).days
                    
                    if days_left == 0:
                        self.add_notification(
                            "ippcu_urgent",
                            f"🚨 СРОЧНО: ИППСУ {last} {first} истекает сегодня!",
                            "error"
                        )
                    elif 0 < days_left <= 7:
                        self.add_notification(
                            "ippcu_warning",
                            f"⚠️ ИППСУ {last} {first} истекает через {days_left} дн.",
                            "warning"
                        )
                    elif 7 < days_left <= 30:
                        self.add_notification(
                            "ippcu_info",
                            f"ℹ️ ИППСУ {last} {first} истекает через {days_left} дн.",
                            "info"
                        )
                    elif days_left < 0:
                        self.add_notification(
                            "ippcu_expired",
                            f"❌ ПРОСРОЧЕНО: ИППСУ {last} {first} ({abs(days_left)} дн. назад)",
                            "error"
                        )
                except ValueError:
                    continue
                    
        except Exception as e:
            print(f"Ошибка проверки ИППСУ: {e}")
    
    def check_empty_contracts(self):
        """Проверка клиентов без договоров"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.cursor()
                cur.execute("""
                    SELECT last_name, first_name 
                    FROM clients 
                    WHERE contract_number IS NULL OR contract_number = '' OR contract_number = 'не указан'
                """)
                
                empty_contracts = cur.fetchall()
            
            if empty_contracts:
                client_list = ", ".join([f"{last} {first}" for last, first in empty_contracts[:3]])
                if len(empty_contracts) > 3:
                    client_list += f" и ещё {len(empty_contracts) - 3}"
                
                self.add_notification(
                    "empty_contracts",
                    f"📄 Отсутствуют номера договоров у {len(empty_contracts)} клиентов: {client_list}",
                    "warning"
                )
                
        except Exception as e:
            print(f"Ошибка проверки договоров: {e}")
    
    def check_upcoming_reviews(self):
        """Проверка предстоящих пересмотров ИППСУ"""
        try:
            today = datetime.today().date()
            next_month = today + timedelta(days=30)
            
            with sqlite3.connect(self.db_path) as conn:
                cur = conn.cursor()
                cur.execute("""
                    SELECT last_name, first_name, ippcu_start 
                    FROM clients 
                    WHERE ippcu_start IS NOT NULL AND ippcu_start != ''
                """)
                
                clients = cur.fetchall()
            
            for last, first, ippcu_start in clients:
                try:
                    if not ippcu_start:
                        continue
                        
                    start_date = datetime.strptime(ippcu_start, "%Y-%m-%d").date()
                    # Предполагаем, что пересмотр нужен через 6 месяцев
                    review_date = start_date + timedelta(days=180)
                    days_until_review = (review_date - today).days
                    
                    if 0 <= days_until_review <= 30:
                        level = "warning" if days_until_review <= 7 else "info"
                        self.add_notification(
                            "review",
                            f"📋 Требуется пересмотр ИППСУ для {last} {first} через {days_until_review} дн.",
                            level
                        )
                except ValueError:
                    continue
                    
        except Exception as e:
            print(f"Ошибка проверки пересмотров: {e}")
    
    def add_notification(self, category, message, level="info"):
        """Добавление уведомления с предотвращением дубликатов"""
        # Проверяем на дубликаты за последние 24 часа
        yesterday = datetime.now() - timedelta(days=1)
        recent_duplicate = any(
            n['category'] == category and 
            n['message'] == message and 
            n['timestamp'] > yesterday and
            not n['read']
            for n in self.notifications
        )
        
        if not recent_duplicate:
            self.notifications.append({
                "id": len(self.notifications) + 1,
                "timestamp": datetime.now(),
                "category": category,
                "message": message,
                "level": level,
                "read": False
            })
    
    def get_unread_count(self):
        """Получить количество непрочитанных уведомлений"""
        return sum(1 for n in self.notifications if not n['read'])
    
    def get_notifications_by_priority(self, unread_only=False):
        """Получить уведомления, отсортированные по приоритету"""
        priority_order = {"error": 0, "warning": 1, "info": 2}
        
        notifications = self.notifications
        if unread_only:
            notifications = [n for n in notifications if not n['read']]
            
        return sorted(notifications, 
                     key=lambda x: (priority_order.get(x['level'], 3), x['timestamp']), 
                     reverse=True)
    
    def mark_as_read(self, notification_id):
        """Пометить уведомление как прочитанное"""
        for notification in self.notifications:
            if notification['id'] == notification_id:
                notification['read'] = True
                break
    
    def mark_all_read(self):
        """Пометить все уведомления как прочитанные"""
        for notification in self.notifications:
            notification['read'] = True
    
    def clear_old_notifications(self, days=7):
        """Очистить старые уведомления"""
        cutoff_date = datetime.now() - timedelta(days=days)
        self.notifications = [
            n for n in self.notifications 
            if n['timestamp'] > cutoff_date or not n['read']
        ]
    
    def show_notification_window(self):
        """Показать окно уведомлений"""
        if not self.is_initialized:
            messagebox.showinfo("Информация", "Система уведомлений не инициализирована")
            return
            
        NotificationWindow(self)

class NotificationWindow:
    def __init__(self, notification_system):
        self.notification_system = notification_system
        self.create_window()
    
    def create_window(self):
        """Создание окна уведомлений"""
        self.window = tk.Toplevel(root)
        self.window.title("🔔 Уведомления")
        self.window.geometry("600x500")
        self.window.configure(bg=ModernStyle.COLORS['background'])
        self.window.minsize(500, 400)
        
        # Заголовок с количеством уведомлений
        header = tk.Frame(self.window, bg=ModernStyle.COLORS['primary'], height=60)
        header.pack(fill='x', padx=0, pady=0)
        header.pack_propagate(False)
        
        unread_count = self.notification_system.get_unread_count()
        title_text = f"🔔 Уведомления ({unread_count} непрочитанных)"
        
        tk.Label(header, text=title_text, 
                bg=ModernStyle.COLORS['primary'],
                fg='white',
                font=ModernStyle.FONTS['h2']).pack(pady=15)
        
        # Основной контент
        main_frame = tk.Frame(self.window, bg=ModernStyle.COLORS['background'])
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Создаем фрейм для списка уведомлений с прокруткой
        list_frame = tk.Frame(main_frame, bg=ModernStyle.COLORS['background'])
        list_frame.pack(fill='both', expand=True)
        
        # Прокрутка
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.notification_canvas = tk.Canvas(
            list_frame, 
            bg=ModernStyle.COLORS['surface'],
            yscrollcommand=scrollbar.set,
            highlightthickness=0
        )
        self.notification_canvas.pack(side='left', fill='both', expand=True)
        scrollbar.config(command=self.notification_canvas.yview)
        
        # Фрейм для уведомлений внутри canvas
        self.notifications_frame = tk.Frame(self.notification_canvas, bg=ModernStyle.COLORS['surface'])
        self.canvas_window = self.notification_canvas.create_window(
            (0, 0), window=self.notifications_frame, anchor='nw', width=self.notification_canvas.winfo_reqwidth()
        )
        
        # Кнопки управления
        button_frame = tk.Frame(main_frame, bg=ModernStyle.COLORS['background'])
        button_frame.pack(fill='x', pady=(10, 0))
        
        ttk.Button(button_frame, text="📁 Пометить все как прочитанные", 
                  style='Primary.TButton',
                  command=self.mark_all_read).pack(side='left', padx=(0, 10))
        
        ttk.Button(button_frame, text="🗑️ Очистить старые", 
                  style='Secondary.TButton',
                  command=self.clear_old).pack(side='left', padx=(0, 10))
        
        ttk.Button(button_frame, text="🔄 Обновить", 
                  style='Secondary.TButton',
                  command=self.refresh).pack(side='left')
        
        ttk.Button(button_frame, text="✖️ Закрыть", 
                  style='Secondary.TButton',
                  command=self.window.destroy).pack(side='right')
        
        # Привязки событий
        self.notifications_frame.bind('<Configure>', self.on_frame_configure)
        self.notification_canvas.bind('<Configure>', self.on_canvas_configure)
        
        self.refresh()
    
    def on_frame_configure(self, event):
        """Обновить scrollregion при изменении размера фрейма"""
        self.notification_canvas.configure(scrollregion=self.notification_canvas.bbox("all"))
    
    def on_canvas_configure(self, event):
        """Обновить ширину внутреннего фрейма при изменении размера canvas"""
        self.notification_canvas.itemconfig(self.canvas_window, width=event.width)
    
    def create_notification_widget(self, parent, notification):
        """Создать виджет для одного уведомления"""
        frame = tk.Frame(parent, bg=ModernStyle.COLORS['surface'], relief='solid', bd=1, padx=10, pady=8)
        frame.pack(fill='x', pady=2)
        
        # Иконки в зависимости от уровня и статуса
        level_icons = {
            'error': '❌',
            'warning': '⚠️',
            'info': 'ℹ️'
        }
        
        status_icon = '✅' if notification['read'] else '🔔'
        level_icon = level_icons.get(notification['level'], '📌')
        
        # Верхняя строка: иконки и время
        top_frame = tk.Frame(frame, bg=ModernStyle.COLORS['surface'])
        top_frame.pack(fill='x')
        
        tk.Label(top_frame, text=f"{status_icon} {level_icon}", 
                bg=ModernStyle.COLORS['surface'],
                fg=ModernStyle.COLORS['text_secondary'],
                font=ModernStyle.FONTS['small']).pack(side='left')
        
        time_str = notification['timestamp'].strftime("%d.%m.%Y %H:%M")
        tk.Label(top_frame, text=time_str,
                bg=ModernStyle.COLORS['surface'],
                fg=ModernStyle.COLORS['text_secondary'],
                font=ModernStyle.FONTS['small']).pack(side='right')
        
        # Текст уведомления
        message_label = tk.Label(frame, text=notification['message'],
                               bg=ModernStyle.COLORS['surface'],
                               fg=ModernStyle.COLORS['text_primary'],
                               font=ModernStyle.FONTS['body'],
                               justify='left',
                               wraplength=550)
        message_label.pack(fill='x', pady=(5, 0))
        
        # Кнопка пометить как прочитанное (только для непрочитанных)
        if not notification['read']:
            def mark_read():
                self.notification_system.mark_as_read(notification['id'])
                self.refresh()
            
            btn_frame = tk.Frame(frame, bg=ModernStyle.COLORS['surface'])
            btn_frame.pack(fill='x', pady=(5, 0))
            
            ttk.Button(btn_frame, text="Отметить как прочитанное",
                      style='Secondary.TButton',
                      command=mark_read).pack(side='right')
        
        return frame
    
    def refresh(self):
        """Обновить список уведомлений"""
        # Очищаем старые виджеты
        for widget in self.notifications_frame.winfo_children():
            widget.destroy()
        
        # Получаем уведомления, отсортированные по приоритету
        notifications = self.notification_system.get_notifications_by_priority()
        
        if not notifications:
            # Сообщение об отсутствии уведомлений
            empty_frame = tk.Frame(self.notifications_frame, bg=ModernStyle.COLORS['surface'], height=100)
            empty_frame.pack(fill='both', expand=True, pady=20)
            empty_frame.pack_propagate(False)
            
            tk.Label(empty_frame, text="🎉 Нет уведомлений",
                    bg=ModernStyle.COLORS['surface'],
                    fg=ModernStyle.COLORS['text_secondary'],
                    font=ModernStyle.FONTS['h3']).pack(expand=True)
            
            tk.Label(empty_frame, text="Все задачи выполнены!",
                    bg=ModernStyle.COLORS['surface'],
                    fg=ModernStyle.COLORS['text_secondary'],
                    font=ModernStyle.FONTS['body']).pack()
        else:
            # Создаем виджеты для каждого уведомления
            for notification in notifications:
                self.create_notification_widget(self.notifications_frame, notification)
    
    def mark_all_read(self):
        """Пометить все как прочитанные"""
        self.notification_system.mark_all_read()
        self.refresh()
    
    def clear_old(self):
        """Очистить старые уведомления"""
        self.notification_system.clear_old_notifications()
        self.refresh()

# Глобальный экземпляр системы уведомлений
notification_system = NotificationSystem(DB_NAME)

# ================== СОВРЕМЕННЫЙ СТИЛЬ ==================
class ModernStyle:
    COLORS = {
        'primary': '#2E86AB',
        'primary_dark': '#1A5A7A',
        'secondary': '#A23B72',
        'accent': '#F18F01',
        'success': '#4CAF50',
        'warning': '#FF9800',
        'error': '#F44336',
        'background': '#F8F9FA',
        'surface': '#FFFFFF',
        'text_primary': '#212529',
        'text_secondary': '#6C757D',
        'border': '#DEE2E6'
    }
    
    FONTS = {
        'h1': ('Segoe UI', 20, 'bold'),
        'h2': ('Segoe UI', 16, 'bold'),
        'h3': ('Segoe UI', 14, 'bold'),
        'body': ('Segoe UI', 11),
        'small': ('Segoe UI', 10),
        'button': ('Segoe UI', 11, 'bold')
    }

def setup_modern_style():
    """Настройка современного стиля"""
    style = ttk.Style()
    
    try:
        style.theme_use('vista')
    except:
        try:
            style.theme_use('clam')
        except:
            pass
    
    # Настраиваем стили
    style.configure('Modern.TFrame', background=ModernStyle.COLORS['background'])
    style.configure('Modern.TLabel', background=ModernStyle.COLORS['background'], 
                   foreground=ModernStyle.COLORS['text_primary'], font=ModernStyle.FONTS['body'])
    style.configure('Primary.TButton', background=ModernStyle.COLORS['primary'], 
                   foreground='white', font=ModernStyle.FONTS['button'], borderwidth=0)
    style.configure('Secondary.TButton', background=ModernStyle.COLORS['surface'], 
                   foreground=ModernStyle.COLORS['primary'], font=ModernStyle.FONTS['button'])
    
    style.map('Primary.TButton',
              background=[('active', ModernStyle.COLORS['primary_dark']),
                         ('pressed', ModernStyle.COLORS['primary_dark'])])
    
    style.map('Secondary.TButton',
              background=[('active', ModernStyle.COLORS['border']),
                         ('pressed', ModernStyle.COLORS['border'])])
    
    # Стиль для Treeview
    style.configure('Modern.Treeview', 
                   background=ModernStyle.COLORS['surface'],
                   fieldbackground=ModernStyle.COLORS['surface'],
                   foreground=ModernStyle.COLORS['text_primary'],
                   font=ModernStyle.FONTS['body'],
                   rowheight=25)
    
    style.configure('Modern.Treeview.Heading', 
                   background=ModernStyle.COLORS['primary'],
                   foreground='white',
                   font=ModernStyle.FONTS['button'],
                   relief='flat')
    
    style.map('Modern.Treeview', 
              background=[('selected', ModernStyle.COLORS['primary'])],
              foreground=[('selected', 'white')])

# ----------------------
# --- Утилиты ФИО ------
# ----------------------
def split_fio(fio: str):
    if not fio:
        return "", "", ""
    parts = fio.strip().split()
    if len(parts) == 1:
        return parts[0], "", ""
    if len(parts) == 2:
        return parts[0], parts[1], ""
    last = parts[0]
    first = parts[1]
    middle = " ".join(parts[2:])
    return last, first, middle

def join_fio(last, first, middle):
    parts = [p for p in (last or "", first or "", middle or "") if p and p.strip()]
    return " ".join(parts)

# ================== База данных ==================
def init_db():
    """Создаёт новую схему или мигрирует старую с обработкой блокировок"""
    max_retries = 5
    retry_delay = 0.1
    
    for attempt in range(max_retries):
        try:
            with sqlite3.connect(DB_NAME, timeout=10.0) as conn:
                cur = conn.cursor()
                
                # Проверяем существование таблицы clients
                cur.execute("""
                    SELECT name FROM sqlite_master 
                    WHERE type='table' AND name='clients'
                """)
                table_exists = cur.fetchone() is not None
                
                if not table_exists:
                    print("Создаем таблицу clients...")
                    cur.execute(
                        """
                        CREATE TABLE IF NOT EXISTS clients (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            last_name TEXT NOT NULL,
                            first_name TEXT NOT NULL,
                            middle_name TEXT,
                            dob TEXT NOT NULL,
                            phone TEXT,
                            contract_number TEXT,
                            ippcu_start TEXT,
                            ippcu_end TEXT,
                            group_name TEXT,
                            UNIQUE(last_name, first_name, middle_name, dob)
                        )
                        """
                    )
                    conn.commit()
                    print("Таблица clients создана успешно")
                    return

                # Проверяем структуру существующей таблицы
                cur.execute("PRAGMA table_info(clients)")
                cols = [r[1] for r in cur.fetchall()]

                if "fio" in cols and "last_name" not in cols:
                    print("Мигрируем старую схему...")
                    cur.execute(
                        """
                        CREATE TABLE IF NOT EXISTS clients_new (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            last_name TEXT NOT NULL,
                            first_name TEXT NOT NULL,
                            middle_name TEXT,
                            dob TEXT NOT NULL,
                            phone TEXT,
                            contract_number TEXT,
                            ippcu_start TEXT,
                            ippcu_end TEXT,
                            group_name TEXT,
                            UNIQUE(last_name, first_name, middle_name, dob)
                        )
                        """
                    )
                    cur.execute("SELECT id, fio, dob, phone, contract_number, ippcu_start, ippcu_end, group_name FROM clients")
                    rows = cur.fetchall()
                    for r in rows:
                        _, fio, dob, phone, contract_number, ippcu_start, ippcu_end, group_name = r
                        last, first, middle = split_fio(fio or "")
                        dob_val = dob or ""
                        try:
                            cur.execute(
                                """
                                INSERT OR IGNORE INTO clients_new
                                (id, last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name)
                                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                """,
                                (None, last, first, middle, dob_val, phone, contract_number, ippcu_start, ippcu_end, group_name)
                            )
                        except Exception:
                            cur.execute(
                                "INSERT OR IGNORE INTO clients_new (last_name, first_name, middle_name, dob) VALUES (?, ?, ?, ?)",
                                (last or "", first or "", middle or "", dob_val)
                            )
                    cur.execute("DROP TABLE clients")
                    cur.execute("ALTER TABLE clients_new RENAME TO clients")
                    conn.commit()
                    print("Миграция завершена успешно")
                    return

                # Добавляем отсутствующие колонки если нужно
                try:
                    if "last_name" not in cols:
                        cur.execute("ALTER TABLE clients ADD COLUMN last_name TEXT")
                    if "first_name" not in cols:
                        cur.execute("ALTER TABLE clients ADD COLUMN first_name TEXT")
                    if "middle_name" not in cols:
                        cur.execute("ALTER TABLE clients ADD COLUMN middle_name TEXT")
                    conn.commit()
                except Exception as e:
                    print(f"Ошибка при добавлении колонок: {e}")

                print("База данных инициализирована успешно")
                break
                
        except sqlite3.OperationalError as e:
            if "locked" in str(e) and attempt < max_retries - 1:
                print(f"База данных заблокирована, повторная попытка {attempt + 1}/{max_retries}...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                print(f"❌ Ошибка инициализации БД: {e}")
                raise e

def add_client(last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group):
    """Добавление с проверкой дублей (по ФИО+дата рождения, без учёта регистра)."""
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        middle_name = middle_name or ""
        dob_val = dob or ""

        cur.execute(
            """
            SELECT id FROM clients
            WHERE lower(last_name)=lower(?) AND lower(first_name)=lower(?) AND lower(COALESCE(middle_name,''))=lower(?) AND dob=?
            """,
            (last_name, first_name, middle_name, dob_val)
        )
        if cur.fetchone():
            raise ValueError(f"Клиент '{join_fio(last_name, first_name, middle_name)}' с датой рождения {dob_val} уже есть в базе.")

        cur.execute(
            """
            INSERT INTO clients (last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (last_name, first_name, middle_name, dob_val, phone, contract_number, ippcu_start, ippcu_end, group),
        )
        conn.commit()

def get_all_clients(limit=200):
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name
            FROM clients
            ORDER BY lower(last_name), lower(first_name)
            LIMIT ?
            """,
            (limit,),
        )
        return cur.fetchall()

def search_clients(query="", date_from=None, date_to=None, limit=200):
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        q = (query or "").strip().lower()
        like = f"%{q}%"

        sql = """
            SELECT id, last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name
            FROM clients
            WHERE ( lower(last_name) LIKE ? OR lower(first_name) LIKE ? OR lower(COALESCE(middle_name,'')) LIKE ?
                   OR lower(contract_number) LIKE ? OR lower(phone) LIKE ? OR lower(COALESCE(group_name,'')) LIKE ? )
        """
        params = [like, like, like, like, like, like]

        if date_from:
            sql += " AND DATE(ippcu_end) >= DATE(?) "
            params.append(date_from)
        if date_to:
            sql += " AND DATE(ippcu_end) <= DATE(?) "
            params.append(date_to)

        sql += " ORDER BY lower(last_name), lower(first_name) LIMIT ?"
        params.append(limit)

        cur.execute(sql, params)
        return cur.fetchall()

def update_client(cid, last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group):
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        cur.execute(
            """
            UPDATE clients
            SET last_name=?, first_name=?, middle_name=?, dob=?, phone=?, contract_number=?, ippcu_start=?, ippcu_end=?, group_name=?
            WHERE id=?
            """,
            (last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group, cid),
        )
        conn.commit()

def delete_client(cid):
    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        cur.execute("DELETE FROM clients WHERE id=?", (cid,))
        conn.commit()

# ================== Google Sheets ==================
def get_gsheet(sheet_id, sheet_name="Лист1"):
    scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
    creds_json = os.getenv("GOOGLE_CREDENTIALS")

    if not creds_json:
        if getattr(sys, "frozen", False):
            exe_dir = os.path.dirname(sys.executable)
        else:
            exe_dir = os.path.dirname(os.path.abspath(__file__))
        creds_path = os.path.join(exe_dir, "credentials.json")
        if not os.path.exists(creds_path):
            raise RuntimeError("Не найден GOOGLE_CREDENTIALS и нет файла credentials.json рядом с программой!")
        with open(creds_path, "r", encoding="utf-8") as f:
            creds_json = f.read()

    creds = Credentials.from_service_account_info(json.loads(creds_json), scopes=scopes)
    client = gspread.authorize(creds)
    sheet = client.open_by_key(sheet_id).worksheet(sheet_name)
    return sheet

def import_from_gsheet():
    try:
        sheet = get_gsheet(SHEET_ID)
        data = sheet.get_all_records()

        added = 0
        for row in data:
            fio_raw = row.get("ФИО", "") or ""
            last, first, middle = split_fio(fio_raw)
            dob = row.get("Дата рождения", "") or ""
            phone = row.get("Телефон", "") or ""
            contract = row.get("Номер договора", "") or ""
            ippcu_start = row.get("Дата начала ИППСУ", "") or ""
            ippcu_end = row.get("Дата окончания ИППСУ", "") or ""
            group = row.get("Группа", "") or ""
            try:
                add_client(last, first, middle, dob, phone, contract, ippcu_start, ippcu_end, group)
                added += 1
            except ValueError:
                continue
        refresh_tree()
        messagebox.showinfo("Успех", f"Импорт из Google Sheets завершён! Добавлено: {added}")
    except Exception as e:
        traceback.print_exc()
        messagebox.showerror("Ошибка", f"Не удалось импортировать:\n{e}")

# ================== КОМПОНЕНТЫ ИНТЕРФЕЙСА ==================
def create_modern_table(parent):
    """Создание современной таблицы с клиентами"""
    # Контейнер для таблицы с прокруткой
    table_container = tk.Frame(parent, bg=ModernStyle.COLORS['background'])
    table_container.pack(fill='both', expand=True, padx=20, pady=10)
    
    # Прокрутка
    scrollbar = ttk.Scrollbar(table_container)
    scrollbar.pack(side='right', fill='y')
    
    # Таблица
    columns = ("✓", "ID", "Фамилия", "Имя", "Отчество", "Дата рождения", 
               "Телефон", "Номер договора", "Дата начала ИППСУ", 
               "Дата окончания ИППСУ", "Группа")
    
    tree = ttk.Treeview(table_container, columns=columns, show="headings", 
                       style='Modern.Treeview', yscrollcommand=scrollbar.set,
                       height=20)
    tree.pack(side='left', fill='both', expand=True)
    scrollbar.config(command=tree.yview)
    
    # Заголовки колонок
    for col in columns:
        tree.heading(col, text=col)
    
    return tree, table_container
    
def create_modern_header(root):
    """Создание современного заголовка"""
    header_frame = tk.Frame(root, bg=ModernStyle.COLORS['primary'], height=80)
    header_frame.pack(fill='x', padx=0, pady=0)
    
    # Основной заголовок
    title_frame = tk.Frame(header_frame, bg=ModernStyle.COLORS['primary'])
    title_frame.pack(fill='x', padx=20, pady=12)
    
    title_label = tk.Label(title_frame, 
                          text="Отделение дневного пребывания",
                          bg=ModernStyle.COLORS['primary'],
                          fg='white',
                          font=ModernStyle.FONTS['h1'])
    title_label.pack(side='left')
    
    subtitle_label = tk.Label(title_frame,
                             text="Полустационарное обслуживание",
                             bg=ModernStyle.COLORS['primary'],
                             fg='white',
                             font=ModernStyle.FONTS['h3'])
    subtitle_label.pack(side='left', padx=(15, 0))
    
    return header_frame

def create_search_panel(root):
    """Создание панели поиска"""
    search_frame = tk.Frame(root, bg=ModernStyle.COLORS['background'], padx=20, pady=15)
    search_frame.pack(fill='x', padx=0, pady=0)
    
    # Поисковая строка
    search_container = tk.Frame(search_frame, bg=ModernStyle.COLORS['surface'], 
                               relief='solid', bd=1, padx=10, pady=8)
    search_container.pack(fill='x', padx=0, pady=0)
    
    tk.Label(search_container, text="🔍 Поиск клиентов:", 
             bg=ModernStyle.COLORS['surface'],
             fg=ModernStyle.COLORS['text_primary'],
             font=ModernStyle.FONTS['h3']).pack(side='left', padx=(0, 10))
    
    search_entry = tk.Entry(search_container, width=40, font=ModernStyle.FONTS['body'],
                           relief='flat', bg=ModernStyle.COLORS['background'], bd=0)
    search_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))
    
    search_btn = ttk.Button(search_container, text="Найти", style='Primary.TButton',
                           command=lambda: do_search())
    search_btn.pack(side='left', padx=(0, 20))
    
    # Фильтры по датам
    filters_frame = tk.Frame(search_container, bg=ModernStyle.COLORS['surface'])
    filters_frame.pack(side='left')
    
    tk.Label(filters_frame, text="ИППСУ до:", 
             bg=ModernStyle.COLORS['surface'],
             fg=ModernStyle.COLORS['text_secondary'],
             font=ModernStyle.FONTS['small']).pack(side='left', padx=(0, 5))
    
    date_from_entry = DateEntry(filters_frame, width=10, date_pattern="dd.mm.yyyy",
                               font=ModernStyle.FONTS['small'], background=ModernStyle.COLORS['primary'],
                               foreground='white', borderwidth=0)
    date_from_entry.pack(side='left', padx=(0, 10))
    
    tk.Label(filters_frame, text="–", 
             bg=ModernStyle.COLORS['surface'],
             fg=ModernStyle.COLORS['text_secondary'],
             font=ModernStyle.FONTS['small']).pack(side='left', padx=(0, 10))
    
    date_to_entry = DateEntry(filters_frame, width=10, date_pattern="dd.mm.yyyy",
                             font=ModernStyle.FONTS['small'], background=ModernStyle.COLORS['primary'],
                             foreground='white', borderwidth=0)
    date_to_entry.pack(side='left', padx=(0, 10))
    
    filter_btn = ttk.Button(filters_frame, text="Применить", style='Secondary.TButton',
                           command=lambda: do_search())
    filter_btn.pack(side='left')
    
    return search_entry, date_from_entry, date_to_entry, search_frame

def create_toolbar(root):
    """Создание панели инструментов"""
    toolbar_frame = tk.Frame(root, bg=ModernStyle.COLORS['surface'], padx=20, pady=10)
    toolbar_frame.pack(fill='x', padx=0, pady=0)
    
    buttons = [
        ("➕ Добавить клиента", add_window, 'Primary.TButton', "Ctrl+N"),
        ("✏️ Редактировать", edit_client, 'Secondary.TButton', "Ctrl+E"),
        ("🗑️ Удалить", delete_selected, 'Secondary.TButton', "Delete"),
        ("👁️ Просмотр", lambda: quick_view_wrapper(), 'Secondary.TButton', "Ctrl+Q"),
        ("📥 Импорт", import_from_gsheet, 'Secondary.TButton', "Ctrl+I"),
        ("📄 Экспорт в Word", export_selected_to_word, 'Secondary.TButton', "Ctrl+W"),
        ("📊 Статистика", show_statistics, 'Secondary.TButton', ""),
        ("🔔 Уведомления", show_notifications, 'Secondary.TButton', "F2"),
        ("⚙️ Настройки", settings_window, 'Secondary.TButton', "")
    ]
    
    for text, command, style_name, shortcut in buttons:
        btn = ttk.Button(toolbar_frame, text=text, command=command, style=style_name)
        btn.pack(side='left', padx=(0, 8))
        
        # Сохраняем ссылки на важные кнопки для управления правами
        if text == "➕ Добавить клиента":
            root.add_btn = btn
        elif text == "🗑️ Удалить":
            root.delete_btn = btn
        
        # Добавляем подсказку с горячей клавишей
        if shortcut:
            tooltip_text = f"{text} ({shortcut})"
            create_tooltip(btn, tooltip_text)

    # Кнопка профиля пользователя
    if AUTH_AVAILABLE:
        profile_btn = ttk.Button(toolbar_frame, text="👤 Профиль", 
                               command=show_user_profile, style='Secondary.TButton')
        profile_btn.pack(side='right', padx=(0, 8))

    # Кнопка справки
    help_btn = ttk.Button(toolbar_frame, text="❓ Справка", 
                         command=show_help, style='Secondary.TButton')
    help_btn.pack(side='right')
    create_tooltip(help_btn, "Справка по горячим клавишам (F1)")
    
    return toolbar_frame

def create_status_bar(root):
    """Создание строки статуса"""
    status_frame = tk.Frame(root, bg=ModernStyle.COLORS['primary'], height=30)
    status_frame.pack(fill='x', side='bottom', padx=0, pady=0)
    status_frame.pack_propagate(False)
    
    # Индикатор состояния БД
    db_status_label = tk.Label(status_frame, text="🟢 БД", 
                              bg=ModernStyle.COLORS['primary'],
                              fg='white', font=ModernStyle.FONTS['small'])
    db_status_label.pack(side='left', padx=(10, 0), pady=5)
    
    status_label = tk.Label(status_frame, text="Готово", 
                           bg=ModernStyle.COLORS['primary'],
                           fg='white', font=ModernStyle.FONTS['small'])
    status_label.pack(side='left', padx=10, pady=5)
    
    word_count_label = tk.Label(status_frame, text="Выбрано для Word: 0", 
                               bg=ModernStyle.COLORS['primary'],
                               fg='white', font=ModernStyle.FONTS['small'])
    word_count_label.pack(side='right', padx=10, pady=5)
    
    # Индикатор пользователя
    user_status_label = tk.Label(status_frame, text="Не авторизован", 
                                bg=ModernStyle.COLORS['primary'],
                                fg='white', font=ModernStyle.FONTS['small'])
    user_status_label.pack(side='right', padx=10, pady=5)
    
    root.status_label = status_label
    root.word_count_label = word_count_label
    root.user_status_label = user_status_label
    root.db_status_label = db_status_label
    
    def update_word_count():
        count = sum(1 for row_id in tree.get_children() 
                   if tree.item(row_id, "values")[0] == "X")
        word_count_label.config(text=f"Выбрано для Word: {count}")
    
    def update_db_status():
        try:
            with sqlite3.connect(DB_NAME, timeout=5.0) as conn:
                cur = conn.cursor()
                cur.execute("SELECT 1")
            db_status_label.config(text="🟢 БД")
        except sqlite3.OperationalError:
            db_status_label.config(text="🔴 БД")
        root.after(5000, update_db_status)  # Проверять каждые 5 секунд
    
    root.update_word_count = update_word_count
    root.after(1000, update_db_status)
    return status_frame

# ================== НАСТРОЙКИ ==================
def settings_window():
    """Окно настроек приложения"""
    settings_win = tk.Toplevel(root)
    settings_win.title("Настройки")
    settings_win.geometry("500x400")
    settings_win.configure(bg=ModernStyle.COLORS['background'])
    settings_win.resizable(False, False)
    
    # Заголовок
    header = tk.Frame(settings_win, bg=ModernStyle.COLORS['primary'], height=50)
    header.pack(fill='x', padx=0, pady=0)
    
    tk.Label(header, text="⚙️ Настройки", 
            bg=ModernStyle.COLORS['primary'],
            fg='white',
            font=ModernStyle.FONTS['h2']).pack(pady=10)
    
    # Основное содержимое
    content_frame = tk.Frame(settings_win, bg=ModernStyle.COLORS['background'], padx=20, pady=20)
    content_frame.pack(fill='both', expand=True)
    
    # Путь для экспорта
    export_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    export_frame.pack(fill='x', pady=10)
    
    tk.Label(export_frame, text="Папка для экспорта по умолчанию:",
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_primary'],
            font=ModernStyle.FONTS['body']).pack(anchor='w')
    
    export_path_frame = tk.Frame(export_frame, bg=ModernStyle.COLORS['background'])
    export_path_frame.pack(fill='x', pady=5)
    
    export_path_var = tk.StringVar(value=settings_manager.get('default_export_path'))
    export_entry = tk.Entry(export_path_frame, textvariable=export_path_var, 
                           font=ModernStyle.FONTS['body'], width=40)
    export_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))
    
    def browse_export_path():
        from tkinter import filedialog
        folder = filedialog.askdirectory(initialdir=export_path_var.get())
        if folder:
            export_path_var.set(folder)
    
    ttk.Button(export_path_frame, text="Обзор", 
              style='Secondary.TButton',
              command=browse_export_path).pack(side='right')
    
    # Настройки уведомлений
    notifications_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    notifications_frame.pack(fill='x', pady=10)
    
    show_notifications_var = tk.BooleanVar(value=settings_manager.get('show_notifications', True))
    notifications_check = ttk.Checkbutton(notifications_frame, 
                                        text="Показывать уведомления при запуске",
                                        variable=show_notifications_var,
                                        style='Modern.TCheckbutton')
    notifications_check.pack(anchor='w')
    
    auto_updates_var = tk.BooleanVar(value=settings_manager.get('auto_check_updates', True))
    updates_check = ttk.Checkbutton(notifications_frame,
                                   text="Автоматически проверять обновления",
                                   variable=auto_updates_var,
                                   style='Modern.TCheckbutton')
    updates_check.pack(anchor='w', pady=(5, 0))
    
    # Кнопки сохранения/отмены
    button_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    button_frame.pack(fill='x', pady=20)
    
    def save_settings():
        settings_manager.set('default_export_path', export_path_var.get())
        settings_manager.set('show_notifications', show_notifications_var.get())
        settings_manager.set('auto_check_updates', auto_updates_var.get())
        messagebox.showinfo("Настройки", "Настройки успешно сохранены!")
        settings_win.destroy()
    
    ttk.Button(button_frame, text="Сохранить", 
              style='Primary.TButton',
              command=save_settings).pack(side='right', padx=(10, 0))
    
    ttk.Button(button_frame, text="Отмена", 
              style='Secondary.TButton',
              command=settings_win.destroy).pack(side='right')
    
    # Информация о приложении
    info_frame = tk.Frame(content_frame, bg=ModernStyle.COLORS['background'])
    info_frame.pack(fill='x', pady=20)
    
    tk.Label(info_frame, text="Информация о приложении:",
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_primary'],
            font=ModernStyle.FONTS['h3']).pack(anchor='w')
    
    info_text = f"""
Версия: 1.0
База данных: {DB_NAME}
Папка приложения: {APP_DIR}
    """
    
    tk.Label(info_frame, text=info_text,
            bg=ModernStyle.COLORS['background'],
            fg=ModernStyle.COLORS['text_secondary'],
            font=ModernStyle.FONTS['small'],
            justify='left').pack(anchor='w', pady=5)

# ================== ГОРЯЧИЕ КЛАВИШИ ==================
def setup_search_behavior():
    """Настройка поведения поиска"""
    def on_search_enter(event):
        do_search()
    
    if hasattr(root, 'search_entry') and root.search_entry:
        root.search_entry.bind('<Return>', on_search_enter)

def setup_keyboard_shortcuts():
    """Настройка горячих клавиш"""
    
    # Основные команды
    root.bind('<Control-n>', lambda e: add_window())
    root.bind('<Control-f>', lambda e: root.search_entry.focus())
    root.bind('<Control-s>', lambda e: do_search())
    root.bind('<Delete>', lambda e: delete_selected())
    root.bind('<F5>', lambda e: refresh_tree())
    root.bind('<F1>', lambda e: show_help())
    
    # Навигация
    root.bind('<Control-q>', lambda e: quick_view_wrapper())
    root.bind('<Control-e>', lambda e: edit_client())
    root.bind('<Control-i>', lambda e: import_from_gsheet())
    root.bind('<Control-w>', lambda e: export_selected_to_word())
    
    # Уведомления
    root.bind('<F2>', lambda e: show_notifications())
    
    # Сообщение в статусной строке о горячих клавишах
    show_status_message("Горячие клавиши активированы. Нажмите F1 для справки.")

def quick_view_wrapper():
    """Обертка для быстрого просмотра с горячей клавишей"""
    selected = tree.selection()
    if selected:
        client_id = tree.item(selected[0], "values")[1]
        quick_view(client_id)
    else:
        messagebox.showinfo("Подсказка", "Выберите клиента для быстрого просмотра")

def show_help():
    """Окно справки по горячим клавишам"""
    help_text = """
📋 ГОРЯЧИЕ КЛАВИШИ:

Основные команды:
Ctrl+N - Добавить клиента
Ctrl+F - Перейти в поиск
Ctrl+S - Выполнить поиск
Delete - Удалить выбранного
F5 - Обновить список

Навигация:
Ctrl+Q - Быстрый просмотр
Ctrl+E - Редактировать
Ctrl+I - Импорт из Google Sheets  
Ctrl+W - Экспорт в Word

Уведомления:
F2 - Показать уведомления

Справка:
F1 - Показать эту справку

Управление таблицей:
←/→ - Изменить ширину колонки
Double Click - Автоподбор колонки
Правый клик - Контекстное меню
"""
    
    help_window = tk.Toplevel(root)
    help_window.title("Справка по горячим клавишам")
    help_window.geometry("500x500")
    help_window.configure(bg=ModernStyle.COLORS['background'])
    
    # Заголовок
    header = tk.Frame(help_window, bg=ModernStyle.COLORS['primary'], height=50)
    header.pack(fill='x', padx=0, pady=0)
    
    tk.Label(header, text="⌨️ Горячие клавиши", 
            bg=ModernStyle.COLORS['primary'],
            fg='white',
            font=ModernStyle.FONTS['h2']).pack(pady=10)
    
    # Текст справки
    text_frame = tk.Frame(help_window, bg=ModernStyle.COLORS['background'])
    text_frame.pack(fill='both', expand=True, padx=20, pady=20)
    
    help_text_widget = tk.Text(text_frame, 
                              font=ModernStyle.FONTS['body'],
                              bg=ModernStyle.COLORS['surface'],
                              fg=ModernStyle.COLORS['text_primary'],
                              wrap='word',
                              padx=10,
                              pady=10)
    help_text_widget.pack(fill='both', expand=True)
    
    help_text_widget.insert('1.0', help_text)
    help_text_widget.config(state='disabled')  # Только для чтения
    
    # Кнопка закрытия
    button_frame = tk.Frame(help_window, bg=ModernStyle.COLORS['background'])
    button_frame.pack(fill='x', padx=20, pady=10)
    
    ttk.Button(button_frame, text="Закрыть", 
              style='Primary.TButton',
              command=help_window.destroy).pack(side='right')

# ================== Автоподбор колонок ==================
def auto_resize_columns(tree, max_width=400):
    """Автоподбор ширины колонок с ограничением по максимальной ширине"""
    tree.update_idletasks()
    
    column_priority = {
        "Фамилия": 2, "Имя": 2, "Отчество": 2, 
        "Дата рождения": 1, "Телефон": 1, "Номер договора": 1,
        "Дата начала ИППСУ": 1, "Дата окончания ИППСУ": 1, "Группа": 1,
        "✓": 0, "ID": 0
    }
    
    for col in tree["columns"]:
        header_text = tree.heading(col)["text"]
        header_width = tk.font.Font().measure(header_text) + 30
        
        content_width = header_width
        for item in tree.get_children():
            cell_value = str(tree.set(item, col))
            cell_width = tk.font.Font().measure(cell_value) + 20
            if cell_width > content_width:
                content_width = cell_width
        
        priority = column_priority.get(header_text, 1)
        if priority == 0:
            final_width = min(content_width, 80)
        elif priority == 2:
            final_width = min(content_width, max_width)
        else:
            final_width = min(content_width, 150)
        
        tree.column(col, width=final_width, minwidth=30)

def setup_tree_behavior(tree):
    """Настройка поведения таблицы"""
    def on_header_click(event):
        region = tree.identify("region", event.x, event.y)
        if region == "separator":
            column = tree.identify_column(event.x)
            col_id = column.replace("#", "")
            columns = tree["columns"]
            if col_id.isdigit() and int(col_id) <= len(columns):
                col_name = columns[int(col_id)-1]
                auto_resize_single_column(tree, col_name)
    
    tree.bind("<Double-1>", on_header_click)

def auto_resize_single_column(tree, col_name):
    """Автоподбор ширины для одной колонки"""
    tree.update_idletasks()
    
    header_text = tree.heading(col_name)["text"]
    header_width = tk.font.Font().measure(header_text) + 30
    
    content_width = header_width
    for item in tree.get_children():
        cell_value = str(tree.set(item, col_name))
        cell_width = tk.font.Font().measure(cell_value) + 20
        if cell_width > content_width:
            content_width = cell_width
    
    final_width = min(content_width, 400)
    tree.column(col_name, width=final_width)

def setup_initial_columns(tree):
    """Начальная настройка колонок"""
    tree.column("✓", width=30, minwidth=20, stretch=False)
    tree.column("ID", width=40, minwidth=30, stretch=False)
    tree.column("Фамилия", width=120, minwidth=80)
    tree.column("Имя", width=120, minwidth=80)
    tree.column("Отчество", width=120, minwidth=80)
    tree.column("Дата рождения", width=100, minwidth=80)
    tree.column("Телефон", width=120, minwidth=80)
    tree.column("Номер договора", width=120, minwidth=80)
    tree.column("Дата начала ИППСУ", width=120, minwidth=80)
    tree.column("Дата окончания ИППСУ", width=120, minwidth=80)
    tree.column("Группа", width=100, minwidth=80)

# ================== Контекстное меню ==================
def show_context_menu(event):
    """Показать контекстное меню по правому клику"""
    item = tree.identify_row(event.y)
    if not item:
        return
    
    tree.selection_set(item)
    context_menu = tk.Menu(root, tearoff=0)
    
    values = tree.item(item, "values")
    client_id = values[1]
    last_name = values[2]
    first_name = values[3]
    client_name = f"{last_name} {first_name}"
    
    context_menu.add_command(
        label=f"Редактировать: {client_name} (Ctrl+E)", 
        command=edit_client
    )
    context_menu.add_command(
        label=f"Удалить: {client_name} (Delete)", 
        command=delete_selected
    )
    context_menu.add_separator()
    context_menu.add_command(
        label="Быстрый просмотр (Ctrl+Q)", 
        command=lambda: quick_view(client_id)
    )
    context_menu.add_command(
        label="Скопировать ФИО", 
        command=lambda: copy_to_clipboard(f"{last_name} {first_name} {values[4] or ''}".strip())
    )
    context_menu.add_command(
        label="Скопировать телефон", 
        command=lambda: copy_to_clipboard(values[6] or "")
    )
    context_menu.add_separator()
    context_menu.add_command(
        label="Добавить в список Word", 
        command=lambda: add_to_word_list(item)
    )
    context_menu.add_separator()
    context_menu.add_command(
        label="Справка по горячим клавишам (F1)", 
        command=show_help
    )
    
    try:
        context_menu.tk_popup(event.x_root, event.y_root)
    finally:
        context_menu.grab_release()

# ================== UI ФУНКЦИИ ==================
def refresh_tree(results=None):
    # очищаем таблицу
    for row in tree.get_children():
        tree.delete(row)

    # если нет результатов — берём все записи
    if results is None:
        results = get_all_clients(limit=200)

    today = datetime.today().date()
    soon = today + timedelta(days=30)

    for row in results:
        cid, last, first, middle, dob, phone, contract, ippcu_start, ippcu_end, group = row
        tag = ""
        try:
            if ippcu_end:
                end_date = datetime.strptime(ippcu_end, "%Y-%m-%d").date()
                if end_date < today:
                    tag = "expired"   # срок истёк
                elif end_date <= soon:
                    tag = "soon"      # истекает скоро
                else:
                    tag = "active"    # ещё действует
        except Exception:
            tag = ""

        tree.insert(
            "",
            "end",
            values=(" ", cid, last, first, middle, dob, phone, contract, ippcu_start, ippcu_end, group),
            tags=(tag,)
        )

    # оформление цветом
    tree.tag_configure("expired", background="#F8D7DA")   # красный (просрочен)
    tree.tag_configure("soon", background="#FFF3CD")      # жёлтый (скоро истечёт)
    tree.tag_configure("active", background="#D4EDDA")    # зелёный (активный)

    
    root.after(100, lambda: auto_resize_columns(tree))

def add_window():
    win = tk.Toplevel()
    win.title("Добавить обслуживаемого")
    win.configure(bg=ModernStyle.COLORS['background'])

    fields = [
        ("Фамилия", 0), ("Имя", 1), ("Отчество", 2),
        ("Дата рождения", 3), ("Телефон", 4), ("Номер договора", 5),
        ("Дата начала ИППСУ", 6), ("Дата окончания ИППСУ", 7), ("Группа", 8)
    ]

    entries = {}
    for field, row in fields:
        tk.Label(win, text=field, bg=ModernStyle.COLORS['background'],
                fg=ModernStyle.COLORS['text_primary'], font=ModernStyle.FONTS['body']).grid(row=row, column=0, padx=10, pady=5, sticky="w")
        
        if "Дата" in field:
            entry = DateEntry(win, width=27, date_pattern="dd.mm.yyyy",
                            font=ModernStyle.FONTS['body'])
        else:
            entry = tk.Entry(win, width=30, font=ModernStyle.FONTS['body'])
        
        entry.grid(row=row, column=1, padx=10, pady=5)
        entries[field] = entry

    def save_client():
        last = entries["Фамилия"].get().strip()
        first = entries["Имя"].get().strip()
        middle = entries["Отчество"].get().strip()
        dob = entries["Дата рождения"].get_date().strftime("%Y-%m-%d")
        phone = entries["Телефон"].get().strip()
        contract_number = entries["Номер договора"].get().strip()
        ippcu_start = entries["Дата начала ИППСУ"].get_date().strftime("%Y-%m-%d")
        ippcu_end = entries["Дата окончания ИППСУ"].get_date().strftime("%Y-%m-%d")
        group = entries["Группа"].get().strip()

        if not last or not first or not dob:
            messagebox.showerror("Ошибка", "Поля 'Фамилия', 'Имя' и 'Дата рождения' обязательны!")
            return

        try:
            add_client(last, first, middle, dob, phone, contract_number, ippcu_start, ippcu_end, group)
            refresh_tree()
            win.destroy()
        except ValueError as ve:
            messagebox.showwarning("Дубликат", str(ve))
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("Ошибка", f"Не удалось добавить:\n{e}")

    save_btn = ttk.Button(win, text="Сохранить", style='Primary.TButton', command=save_client)
    save_btn.grid(row=9, column=0, columnspan=2, pady=10)

def edit_client():
    """Окно редактирования клиента"""
    selected = tree.selection()
    if not selected:
        messagebox.showwarning("Ошибка", "Выберите клиента для редактирования")
        return

    values = tree.item(selected[0], "values")
    cid = values[1]

    last, first, middle = values[2], values[3], values[4]
    dob, phone, contract = values[5], values[6], values[7]
    ippcu_start, ippcu_end, group = values[8], values[9], values[10]

    win = tk.Toplevel(root)
    win.title("Редактировать клиента")
    win.configure(bg=ModernStyle.COLORS['background'])

    fields = [
        ("Фамилия", last, 0), ("Имя", first, 1), ("Отчество", middle, 2),
        ("Дата рождения", dob, 3), ("Телефон", phone, 4), ("Номер договора", contract, 5),
        ("Дата начала ИППСУ", ippcu_start, 6), ("Дата окончания ИППСУ", ippcu_end, 7), ("Группа", group, 8)
    ]

    entries = {}
    for field, value, row in fields:
        tk.Label(win, text=field, bg=ModernStyle.COLORS['background'],
                fg=ModernStyle.COLORS['text_primary'], font=ModernStyle.FONTS['body']).grid(row=row, column=0, padx=10, pady=5, sticky="w")
        
        if "Дата" in field:
            entry = tk.Entry(win, width=30, font=ModernStyle.FONTS['body'])
            entry.insert(0, value)
        else:
            entry = tk.Entry(win, width=30, font=ModernStyle.FONTS['body'])
            entry.insert(0, value)
        
        entry.grid(row=row, column=1, padx=10, pady=5)
        entries[field] = entry

    def save_changes():
        update_client(cid,
                      entries["Фамилия"].get(), entries["Имя"].get(), entries["Отчество"].get(),
                      entries["Дата рождения"].get(), entries["Телефон"].get(), entries["Номер договора"].get(),
                      entries["Дата начала ИППСУ"].get(), entries["Дата окончания ИППСУ"].get(), entries["Группа"].get())
        refresh_tree()
        win.destroy()

    save_btn = ttk.Button(win, text="Сохранить", style='Primary.TButton', command=save_changes)
    save_btn.grid(row=9, column=0, columnspan=2, pady=10)

def delete_selected():
    selected = tree.selection()
    if not selected:
        messagebox.showerror("Ошибка", "Выберите клиента для удаления")
        return
    item = tree.item(selected[0])
    cid = item["values"][1]
    if messagebox.askyesno("Удалить", "Точно удалить выбранного клиента?"):
        delete_client(cid)
        refresh_tree()

def do_search():
    query = root.search_entry.get().strip()
    date_from = root.date_from_entry.get_date().strftime("%Y-%m-%d") if root.date_from_entry.get() else None
    date_to = root.date_to_entry.get_date().strftime("%Y-%m-%d") if root.date_to_entry.get() else None

    with sqlite3.connect(DB_NAME) as conn:
        cur = conn.cursor()
        q = (query or "").strip().lower()
        like = f"%{q}%"

        sql = """
            SELECT id, last_name, first_name, middle_name, dob, phone, contract_number, ippcu_start, ippcu_end, group_name
            FROM clients
            WHERE (
                lower(last_name) LIKE ?
                OR lower(first_name) LIKE ?
                OR lower(COALESCE(middle_name,'')) LIKE ?
                OR lower(last_name || ' ' || first_name || ' ' || COALESCE(middle_name,'')) LIKE ?
                OR lower(contract_number) LIKE ?
                OR lower(phone) LIKE ?
                OR lower(COALESCE(group_name,'')) LIKE ?
            )
        """
        params = [like, like, like, like, like, like, like]

        if date_from:
            sql += " AND DATE(ippcu_end) >= DATE(?) "
            params.append(date_from)
        if date_to:
            sql += " AND DATE(ippcu_end) <= DATE(?) "
            params.append(date_to)

        sql += " ORDER BY lower(last_name), lower(first_name) LIMIT ?"
        params.append(200)

        cur.execute(sql, params)
        results = cur.fetchall()

    refresh_tree(results)

def toggle_check(event):
    region = tree.identify("region", event.x, event.y)
    if region != "cell":
        return
    col = tree.identify_column(event.x)
    if col != "#1":
        return

    row_id = tree.identify_row(event.y)
    if not row_id:
        return

    values = list(tree.item(row_id, "values"))
    current = values[0]
    values[0] = "X" if current.strip() == "" else " "
    tree.item(row_id, values=values)
    if hasattr(root, 'update_word_count'):
        root.update_word_count()

# ================== ЧАТ СИСТЕМА ==================
def initialize_chat_system(notebook):
    """Инициализация системы чата"""
    try:
        from chat_manager import ChatManager
        from chat_ui import ChatUI
        from chat_notifications import ChatNotifications
        
        # Инициализация чата
        chat_manager = ChatManager()
        chat_notifications = ChatNotifications(chat_manager)
        
        # Создание UI чата
        chat_ui = ChatUI(notebook, chat_manager, ModernStyle.COLORS, ModernStyle.FONTS)
        chat_frame = chat_ui.get_widget()
        notebook.add(chat_frame, text="💬 Чат сотрудников")
        
        # Сохраняем ссылки для доступа из других функций
        root.chat_manager = chat_manager
        root.chat_ui = chat_ui
        root.chat_notifications = chat_notifications
        
        # Функция периодического обновления чата
        def update_chat_periodically():
            if hasattr(root, 'chat_ui') and root.chat_ui:
                try:
                    root.chat_ui.refresh_chat()
                    root.chat_ui.update_unread_count()
                except Exception as e:
                    print(f"Ошибка обновления чата: {e}")
            root.after(30000, update_chat_periodically)
        
        root.after(5000, update_chat_periodically)
        root.after(4000, lambda: chat_manager.set_user_online(True))
        
        print("✅ Модуль чата инициализирован")
        return True
        
    except ImportError as e:
        print(f"❌ Модули чата не найдены: {e}")
        return False
    except Exception as e:
        print(f"❌ Ошибка инициализации чата: {e}")
        return False

def create_chat_stub(notebook):
    """Создание заглушки для чата"""
    chat_stub_frame = tk.Frame(notebook, bg=ModernStyle.COLORS['background'])
    notebook.add(chat_stub_frame, text="💬 Чат (недоступен)")
    
    stub_label = tk.Label(chat_stub_frame, 
                        text="Модуль чата не установлен\n\nДля использования чата установите необходимые зависимости",
                        bg=ModernStyle.COLORS['background'],
                        fg=ModernStyle.COLORS['text_secondary'],
                        font=ModernStyle.FONTS['h3'],
                        justify='center')
    stub_label.pack(expand=True, fill='both', padx=20, pady=20)

# ================== MAIN ==================
def main():
    global root, tree, auth_manager
    
    # Инициализация главного окна
    root = tk.Tk()
    root.title("Отделение дневного пребывания - Авторизация")
    root.geometry("1400x900")
    root.configure(bg=ModernStyle.COLORS['background'])
    
    # Показываем сообщение о загрузке
    loading_label = tk.Label(root, text="Загрузка приложения...", 
                            bg=ModernStyle.COLORS['background'],
                            font=ModernStyle.FONTS['h2'])
    loading_label.pack(expand=True)
    root.update()
    
    try:
        # Инициализация системы аутентификации
        setup_auth_system()
        
        # Инициализация БД
        init_db()
        
        loading_label.destroy()
        
        # Показываем окно входа или основное приложение
        if AUTH_AVAILABLE and auth_manager and (not getattr(auth_manager, 'current_user', None) or not getattr(auth_manager, 'remember_me', False)):
            show_login_window()
        else:
            initialize_main_application()
            
        root.mainloop()
        
    except Exception as e:
        messagebox.showerror("Критическая ошибка", 
                           f"Не удалось запустить приложение:\n{e}")
        root.destroy()

def initialize_main_application():
    """Инициализация основного приложения после авторизации"""
    print("DEBUG: Starting initialize_main_application")
    
    try:
        # Обновляем заголовок окна
        if AUTH_AVAILABLE and auth_manager and auth_manager.current_user:
            root.title(f"Отделение дневного пребывания - {auth_manager.get_user_display_name()}")
        else:
            root.title("Отделение дневного пребывания - Демо-режим")
        
        # Инициализация базы данных клиентов
        init_db()
        print("✅ База данных инициализирована")
        
        # Настройка современного стиля
        setup_modern_style()
        
        # Создание Notebook для вкладок
        notebook = ttk.Notebook(root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # === ОСНОВНАЯ ВКЛАДКА - КЛИЕНТЫ ===
        main_frame = tk.Frame(notebook, bg=ModernStyle.COLORS['background'])
        notebook.add(main_frame, text="📋 Клиенты")
        
        # Создание интерфейса в основной вкладке
        header = create_modern_header(main_frame)
        search_entry, date_from_entry, date_to_entry, search_frame = create_search_panel(main_frame)
        toolbar = create_toolbar(main_frame)
        
        # СОЗДАЕМ ТАБЛИЦУ
        global tree
        tree, table_container = create_modern_table(main_frame)
        
        status_bar = create_status_bar(main_frame)
        
        # Сохраняем ссылки на элементы
        root.search_entry = search_entry
        root.date_from_entry = date_from_entry
        root.date_to_entry = date_to_entry
        root.notebook = notebook
        
        # Настройка таблицы
        setup_initial_columns(tree)
        setup_tree_behavior(tree)
        
        # Настройка горячих клавиш
        setup_keyboard_shortcuts()
        setup_search_behavior()
        
        # Привязка событий
        tree.bind("<Button-3>", show_context_menu)
        tree.bind("<Button-1>", toggle_check)
        
        print("DEBUG: Basic UI created, loading data...")
        
        # Загрузка данных
        refresh_tree()
        
        print("✅ Основной интерфейс создан")
        
        # === ВКЛАДКА ЧАТА ===
        def initialize_chat():
            if not initialize_chat_system(notebook):
                create_chat_stub(notebook)
        
        # Инициализируем чат с задержкой
        root.after(1000, initialize_chat)
        
        # === ОТЛОЖЕННЫЕ ОПЕРАЦИИ ===
        
        def load_application_data():
            """Загрузка данных приложения"""
            try:
                # Загрузка данных в таблицу
                refresh_tree()
                print("✅ Таблица клиентов загружена")
                
                # Проверка обновлений
                updater.auto_update()
                print("✅ Проверка обновлений выполнена")
                
            except Exception as e:
                print(f"❌ Ошибка загрузки данных приложения: {e}")
                messagebox.showwarning("Предупреждение", 
                                    f"Некоторые функции могут работать некорректно: {e}")
        
        def initialize_notifications():
            """Инициализация системы уведомлений"""
            try:
                if notification_system.initialize():
                    print("✅ Система уведомлений инициализирована")
                    
                    unread_count = notification_system.get_unread_count()
                    if unread_count > 0:
                        print(f"✅ Найдено {unread_count} непрочитанных уведомлений")
                    else:
                        print("✅ Непрочитанных уведомлений нет")
                else:
                    print("⚠️ Система уведомлений отключена")
            except Exception as e:
                print(f"❌ Ошибка инициализации уведомлений: {e}")
        
        def initialize_security_checks():
            """Инициализация проверок безопасности"""
            try:
                check_expiring_ippcu()
                print("✅ Проверка ИППСУ выполнена")
            except Exception as e:
                print(f"❌ Ошибка проверки ИППСУ: {e}")
        
        def show_welcome_message():
            """Показать приветственное сообщение"""
            if AUTH_AVAILABLE and auth_manager.remember_me:
                show_status_message(f"Автоматический вход: {auth_manager.get_user_display_name()}")
            elif AUTH_AVAILABLE:
                show_status_message(f"Добро пожаловать, {auth_manager.get_user_display_name()}!")
            else:
                show_status_message("Демо-режим: аутентификация отключена")
        
        # Планируем отложенные операции с правильным порядком
        root.after(500, load_application_data)        # Загрузка данных
        root.after(1000, initialize_notifications)    # Уведомления
        root.after(1500, initialize_security_checks)  # Проверки безопасности
        root.after(2000, show_welcome_message)        # Приветственное сообщение
        
        # === ОБРАБОТКА ЗАКРЫТИЯ ПРИЛОЖЕНИЯ ===
        def on_closing():
            """Обработчик закрытия приложения"""
            try:
                # Устанавливаем пользователя оффлайн в чате
                if hasattr(root, 'chat_manager') and root.chat_manager:
                    root.chat_manager.set_user_online(False)
                    print("✅ Пользователь установлен как оффлайн")
                
                # Сохраняем настройки
                settings_manager.save_settings()
                print("✅ Настройки сохранены")
                
                # Очищаем старые уведомления
                if notification_system.is_initialized:
                    notification_system.clear_old_notifications()
                    print("✅ Старые уведомления очищены")
                    
            except Exception as e:
                print(f"⚠️ Ошибка при завершении работы: {e}")
            finally:
                root.destroy()
        
        root.protocol("WM_DELETE_WINDOW", on_closing)
        
        # === СТАТУС ЗАПУСКА ===
        def show_startup_status():
            """Показать статус запуска в статусной строке"""
            if hasattr(root, 'status_label'):
                root.status_label.config(text="Приложение готово к работе")
        
        root.after(3000, show_startup_status)
        
    except Exception as e:
        print(f"❌ Ошибка инициализации приложения: {e}")
        messagebox.showerror("Ошибка", f"Не удалось инициализировать приложение: {e}")

if __name__ == "__main__":
    main()
